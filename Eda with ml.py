import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import io
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, confusion_matrix, classification_report, roc_auc_score, roc_curve
from docx import Document
from docx.shared import Inches
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.tree import DecisionTreeClassifier
from xgboost import XGBClassifier

df = pd.read_csv("diabetes.csv")


cols_with_zeros = ['Glucose', 'BloodPressure', 'SkinThickness', 'Insulin', 'BMI']
df[cols_with_zeros] = df[cols_with_zeros].replace(0, np.nan)
df.fillna(df.median(), inplace=True)

doc = Document()
doc.add_heading("Diabetes Prediction Report (Enhanced)", 0)

doc.add_heading("Dataset Overview", level=1)
buffer = io.StringIO()
df.info(buf=buffer)
doc.add_paragraph(buffer.getvalue())

doc.add_heading("Descriptive Statistics", level=1)
doc.add_paragraph(df.describe().to_string())

doc.add_heading("Missing Value Handling", level=1)
doc.add_paragraph(
    "Zero values in ['Glucose', 'BloodPressure', 'SkinThickness', 'Insulin', 'BMI'] "
    "were treated as missing and replaced using median imputation."
)


plt.figure(figsize=(6, 4))
sns.countplot(x='Outcome', data=df)
plt.title("Diabetes Distribution")
plt.savefig("distribution.png")
plt.close()
doc.add_heading("Diabetes Distribution", level=1)
doc.add_picture("distribution.png", width=Inches(5))


sns.pairplot(df, hue='Outcome')
plt.savefig("pairplot.png")
plt.close()
doc.add_heading("Feature Pairplot", level=1)
doc.add_picture("pairplot.png", width=Inches(6))


X = df.drop('Outcome', axis=1)
y = df['Outcome']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
scaler = StandardScaler()
X_train = scaler.fit_transform(X_train)
X_test = scaler.transform(X_test)


models = {
    "Logistic Regression": LogisticRegression(max_iter=1000),
    "Random Forest": RandomForestClassifier(random_state=42),
    "Support Vector Machine": SVC(probability=True, random_state=42),
    "K-Nearest Neighbors": KNeighborsClassifier(),
    "Decision Tree": DecisionTreeClassifier(random_state=42),
    "XGBoost": XGBClassifier(use_label_encoder=False, eval_metric='logloss', random_state=42)
}
metrics = {}

for name, model in models.items():
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    y_proba = model.predict_proba(X_test)[:, 1]
    
    acc = accuracy_score(y_test, y_pred)
    cm = confusion_matrix(y_test, y_pred)
    cr = classification_report(y_test, y_pred)
    auc = roc_auc_score(y_test, y_proba)
    
    
    metrics[name] = {
        "accuracy": acc,
        "confusion_matrix": cm,
        "classification_report": cr,
        "auc": auc,
        "model": model,
        "y_proba": y_proba
    }

    
    doc.add_heading(f"{name} Evaluation", level=1)
    doc.add_paragraph(f"Accuracy: {acc:.4f}")
    doc.add_paragraph(f"AUC Score: {auc:.4f}")
    doc.add_paragraph("Confusion Matrix:\n" + str(cm))
    doc.add_paragraph("Classification Report:\n" + cr)

    
    plt.figure(figsize=(5, 4))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues')
    plt.title(f"{name} - Confusion Matrix")
    plt.xlabel("Predicted")
    plt.ylabel("Actual")
    plt.tight_layout()
    cm_path = f"{name.replace(' ', '_').lower()}_cm.png"
    plt.savefig(cm_path)
    plt.close()
    doc.add_picture(cm_path, width=Inches(5))

    
    fpr, tpr, thresholds = roc_curve(y_test, y_proba)
    plt.figure()
    plt.plot(fpr, tpr, label=f"{name} (AUC = {auc:.2f})")
    plt.plot([0, 1], [0, 1], 'k--')
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title(f"{name} - ROC Curve")
    plt.legend()
    roc_path = f"{name.replace(' ', '_').lower()}_roc.png"
    plt.savefig(roc_path)
    plt.close()
    doc.add_picture(roc_path, width=Inches(5))

lr_importance = models["Logistic Regression"].coef_[0]
features = df.columns[:-1]
plt.figure(figsize=(8, 5))
plt.barh(features, lr_importance)
plt.title("Feature Importance - Logistic Regression")
plt.xlabel("Coefficient Value")
plt.tight_layout()
plt.savefig("feature_importance_logreg.png")
plt.close()
doc.add_heading("Feature Importance - Logistic Regression", level=1)
doc.add_picture("feature_importance_logreg.png", width=Inches(5.5))

doc.save("Diabetes_Prediction_Report_Enhanced.docx")
print(" Report saved as: Diabetes_Prediction_Report_Enhanced.docx")

import joblib
joblib.dump(model, "diabetes_model.pkl")
joblib.dump(scaler, "scaler.pkl")
